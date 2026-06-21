# -*- coding: utf-8 -*-
"""Doc CV cua ban (1 file duy nhat) -> text. Ho tro .pdf / .docx / .txt."""
import os
import re
import glob

SAMPLE_FALLBACK = "CV_NguyenMinhAnh_Marketing.docx"  # CV mau de test khi chua co CV that


def _maybe_despace(text: str) -> str:
    """
    Mot so PDF dung letter-spacing -> text bi gian "S o c i a l  C o n t e n t".
    Neu phat hien (nhieu token 1 ky tu), gom lai: trong 1 dong, tu cach nhau bang
    2+ space, chu cai trong tu cach nhau bang 1 space.
    """
    tokens = text.split()
    if not tokens:
        return text
    singles = sum(1 for t in tokens if len(t) == 1)
    if singles / len(tokens) < 0.4:
        return text  # text binh thuong, khong dung
    out_lines = []
    for line in text.splitlines():
        words = re.split(r"\s{2,}", line.strip())
        words = [re.sub(r"(?<=\S) (?=\S)", "", w) for w in words]
        out_lines.append(" ".join(w for w in words if w))
    return "\n".join(out_lines)


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # lay them text trong bang (ky nang thuong nam o bang)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(t for t in parts if t and t.strip()).strip()


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def find_cv_path() -> str:
    """
    Thu tu uu tien:
      1. Bien moi truong CV_PATH
      2. File dau tien trong thu muc cv/  (*.pdf, *.docx, *.txt)
      3. CV mau trong repo (de test luong)
    """
    env_path = os.getenv("CV_PATH", "").strip()
    if env_path and os.path.exists(env_path):
        return env_path

    for ext in ("pdf", "docx", "txt"):
        hits = sorted(glob.glob(os.path.join("cv", f"*.{ext}")))
        if hits:
            return hits[0]

    if os.path.exists(SAMPLE_FALLBACK):
        print(f"[CV] ! Chua co CV that trong cv/ — dung tam CV mau: {SAMPLE_FALLBACK}")
        return SAMPLE_FALLBACK

    raise FileNotFoundError(
        "Khong tim thay CV. Hay dat file CV (.pdf/.docx/.txt) vao thu muc cv/ "
        "hoac set bien moi truong CV_PATH."
    )


def load_cv_text(path: str = None) -> tuple[str, str]:
    """Tra ve (cv_text, path_da_dung)."""
    path = path or find_cv_path()
    ext = path.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        text = _read_pdf(path)
    elif ext == "docx":
        text = _read_docx(path)
    else:
        text = _read_txt(path)

    text = _maybe_despace(text)

    if not text or len(text) < 30:
        raise ValueError(f"Doc CV ra qua it noi dung ({len(text)} ky tu): {path}")
    print(f"[CV] Da doc CV: {path}  ({len(text)} ky tu)")
    return text, path


if __name__ == "__main__":
    from dotenv import load_dotenv
    load_dotenv()
    txt, p = load_cv_text()
    print(f"--- Preview ({p}) ---")
    print(txt[:600])
