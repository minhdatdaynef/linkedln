"""
Tao CV Marketing mau bang python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, side: str, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"),   "single")
    border.set(qn("w:sz"),    sz)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)
    tcBorders.append(border)

def hr(doc, color="0A66C2", space_before=6, space_after=6):
    """Horizontal divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_run(para, text, bold=False, italic=False, size=11, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run

def section_title(doc, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(title.upper())
    run.bold          = True
    run.font.size     = Pt(10.5)
    run.font.name     = "Arial"
    run.font.color.rgb = RGBColor(0x0A, 0x66, 0xC2)
    hr(doc, color="0A66C2", space_before=0, space_after=6)

def bullet(doc, text: str, indent=0.35):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(indent)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Arial"

def job_header(doc, title, company, period, location="Hà Nội"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    add_run(p, title,   bold=True,  size=11)
    add_run(p, "  |  ",             size=10, color="888888")
    add_run(p, company, bold=False, size=10.5, color="0A66C2")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    add_run(p2, f"📅 {period}  ·  📍 {location}", italic=True, size=9.5, color="888888")

# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Header block ─────────────────────────────────────────────────────────────
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name_para.add_run("NGUYỄN MINH ANH")
name_run.bold          = True
name_run.font.size     = Pt(22)
name_run.font.name     = "Arial"
name_run.font.color.rgb = RGBColor(0x0A, 0x66, 0xC2)
name_para.paragraph_format.space_after = Pt(2)

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_para, "Marketing Executive  ·  Content & Digital Marketing", size=12, color="444444")
title_para.paragraph_format.space_after = Pt(4)

contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_para.paragraph_format.space_after = Pt(2)
add_run(contact_para, "📧 minhanh.marketing@gmail.com  ·  📞 0912 345 678  ·  "
                      "🔗 linkedin.com/in/minhanh-mkt  ·  📍 Hà Nội, Việt Nam",
        size=9.5, color="555555")

hr(doc, color="0A66C2", space_before=4, space_after=8)

# ── Mục tiêu ──────────────────────────────────────────────────────────────────
section_title(doc, "Mục tiêu nghề nghiệp")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_run(p,
    "Marketing Executive với 2+ năm kinh nghiệm về Content Marketing và Digital Advertising. "
    "Mong muốn đóng góp vào chiến lược xây dựng thương hiệu và tăng trưởng doanh thu tại môi trường "
    "năng động, nơi sáng tạo và dữ liệu cùng đồng hành.",
    size=10.5)

# ── Kinh nghiệm ───────────────────────────────────────────────────────────────
section_title(doc, "Kinh nghiệm làm việc")

job_header(doc, "Content Marketing Executive", "Công ty CP Thương mại ABC", "06/2022 – Hiện tại")
for b in [
    "Lên kế hoạch & sản xuất nội dung cho fanpage Facebook 150k followers — tăng engagement rate từ 1.2% lên 3.8% trong 6 tháng",
    "Viết bài blog SEO, tăng organic traffic website từ 5,000 → 22,000 lượt/tháng sau 1 năm",
    "Quản lý lịch đăng bài đa kênh: Facebook, Instagram, TikTok; phối hợp với designer sản xuất banner, infographic, video ngắn",
    "Chạy Facebook Ads ngân sách 30 triệu/tháng — ROAS trung bình đạt 4.2x",
    "Campaign \"Tết Sum Vầy\" 2024: đạt 2.1 triệu reach organic, 1 video TikTok viral 800k views, +12,000 follower mới",
]:
    bullet(doc, b)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

job_header(doc, "Marketing Intern", "Agency Sáng Tạo XYZ", "01/2022 – 05/2022")
for b in [
    "Viết caption, script video TikTok cho 3 nhãn hàng FMCG",
    "Nghiên cứu đối thủ cạnh tranh, tổng hợp báo cáo hàng tuần cho Account Manager",
    "Tham gia tổ chức 2 sự kiện ra mắt sản phẩm offline (200–300 khách mỗi sự kiện)",
]:
    bullet(doc, b)

# ── Học vấn ───────────────────────────────────────────────────────────────────
section_title(doc, "Học vấn")
job_header(doc, "Cử nhân Truyền thông & Quan hệ Công chúng",
           "Đại học Ngoại thương Hà Nội", "2018 – 2022")
p = doc.add_paragraph()
add_run(p, "GPA: 3.4 / 4.0  ·  Tốt nghiệp loại Giỏi", italic=True, size=10, color="555555")
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after  = Pt(2)

# ── Kỹ năng ──────────────────────────────────────────────────────────────────
section_title(doc, "Kỹ năng")

# Two-column skill table (no borders — clean look)
table = doc.add_table(rows=4, cols=2)
table.style = "Table Grid"

skills = [
    ("Content & Copywriting",     "SEO Content, Script Video, Storytelling, Email Marketing"),
    ("Digital Advertising",        "Facebook Ads, Google Ads (cơ bản), Meta Business Suite"),
    ("Tools",                      "Canva, CapCut, Google Analytics, Mailchimp, Trello, Notion"),
    ("Phân tích",                  "Google Analytics, Facebook Insights, báo cáo KPI hàng tháng"),
]

for i, (skill, detail) in enumerate(skills):
    row = table.rows[i]
    for col in row.cells:
        for side in ["top", "bottom", "left", "right"]:
            set_cell_border(col, side, color="E8E8E8", sz="4")

    # Left cell — skill name
    lc = row.cells[0]
    lc.width = Inches(2.1)
    set_cell_bg(lc, "EFF4FB")
    p = lc.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.08)
    run = p.add_run(skill)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"

    # Right cell — detail
    rc = row.cells[1]
    p2 = rc.paragraphs[0]
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after  = Pt(3)
    p2.paragraph_format.left_indent  = Inches(0.08)
    run2 = p2.add_run(detail)
    run2.font.size = Pt(10)
    run2.font.name = "Arial"

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Chứng chỉ ────────────────────────────────────────────────────────────────
section_title(doc, "Chứng chỉ")
for b in [
    "Google Digital Marketing & E-commerce Certificate (2023)",
    "Meta Social Media Marketing Professional Certificate (2022)",
    "HubSpot Content Marketing Certification (2023)",
]:
    bullet(doc, b)

# ── Ngoại ngữ ─────────────────────────────────────────────────────────────────
section_title(doc, "Ngoại ngữ")
p = doc.add_paragraph()
add_run(p, "Tiếng Anh: ", bold=True, size=10.5)
add_run(p, "B2 (IELTS 6.5) — đọc tài liệu chuyên ngành, viết email, thuyết trình tốt", size=10.5)

# ── Hoạt động ngoại khóa ─────────────────────────────────────────────────────
section_title(doc, "Hoạt động & Dự án nổi bật")
for b in [
    "Thành viên Ban tổ chức CLB Marketing FTU — tổ chức hội thảo 500+ người tham dự (2019-2021)",
    "Viết blog cá nhân về Digital Marketing: 2,000 readers/tháng, chủ đề Content Strategy & SEO",
    "Volunteer Content Lead @ TEDxHanoi 2021 — quản lý social media trong suốt sự kiện",
]:
    bullet(doc, b)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"C:\Users\ADMIN\PycharmProjects\linkedin-crawler\CV_NguyenMinhAnh_Marketing.docx"
doc.save(out)
print(f"[OK] Saved: {out}")
